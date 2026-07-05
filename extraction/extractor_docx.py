"""
extractor_docx.py — Text extraction from Microsoft Word DOCX bank statement files.

Some banks, particularly smaller cooperative banks and credit societies, provide
bank statements as Microsoft Word (.docx) documents. These may be formatted as
either plain paragraphs (each line is a transaction) or as Word tables (a grid
with rows and columns, similar to a spreadsheet).

This module reads both formats using the python-docx library:
  1. All paragraphs are read — each paragraph is one line of text.
  2. All tables are read — each cell in each row of each table is extracted.

The combined text is treated identically to digital PDF text and passed to
the column identifier (column_identifier.py), which uses an LLM to determine
which column position holds the date, narration, debit, credit, and balance.

Team: Survey Corps | CIDECODE Hackathon 2026 | CID Karnataka
"""

import logging
from pathlib import Path

import docx  # python-docx library, imported as "docx"

# Set up a logger for this module.
logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts all text from a DOCX bank statement file.

    Uses python-docx to read every paragraph and table cell.
    Bank statements in DOCX format sometimes have transactions
    inside Word tables, so both paragraphs and tables are read.

    The combined text is passed to Component 3 (column identifier)
    exactly like digital PDF text.

    Parameters:
        file_path (str): Absolute path to the DOCX file.

    Returns:
        str: Full extracted text from all paragraphs and tables,
             with each paragraph and each table row on its own line.
             Returns empty string if extraction fails.
    """
    logger.info(
        "extractor_docx.extract_text_from_docx: "
        "Starting extraction from '%s'",
        Path(file_path).name,
    )

    try:
        # Open the DOCX file using python-docx.
        doc = docx.Document(file_path)

        # We collect all text lines into this list, then join them at the end.
        all_lines = []

        # ── Step 1: Extract text from all paragraphs ───────────────────────
        # In Word documents, paragraphs include body text, headings, and
        # individual text blocks. Bank statement data in paragraph form
        # has each transaction as a separate paragraph.
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Skip empty paragraphs (blank lines)
                all_lines.append(text)
                paragraph_count += 1

        logger.info(
            "extractor_docx.extract_text_from_docx: "
            "Extracted %d non-empty paragraphs",
            paragraph_count,
        )

        # ── Step 2: Extract text from all tables ────────────────────────────
        # Many bank DOCX statements use Word tables (like a spreadsheet grid)
        # to organise transaction data into rows and columns.
        # We read each row by combining all cell texts with a tab separator,
        # which mimics the column-separated format that the standardiser expects.
        table_row_count = 0
        for table_index, table in enumerate(doc.tables):
            for row in table.rows:
                # Collect all cell texts in this row
                cell_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_texts.append(cell_text)

                # Join cells with a tab character to create a column-separated line.
                # The standardiser and column identifier will split on whitespace/tabs.
                row_line = "\t".join(cell_texts)
                if row_line.strip():  # Skip completely empty rows
                    all_lines.append(row_line)
                    table_row_count += 1

        logger.info(
            "extractor_docx.extract_text_from_docx: "
            "Extracted %d rows from %d table(s)",
            table_row_count,
            len(doc.tables),
        )

        # Join all lines with newline characters to create one continuous text string.
        combined_text = "\n".join(all_lines)

        logger.info(
            "extractor_docx.extract_text_from_docx: "
            "Extraction complete. Total characters: %d",
            len(combined_text),
        )

        return combined_text

    except FileNotFoundError:
        logger.error(
            "extractor_docx.extract_text_from_docx: File not found: %s",
            file_path,
        )
        return ""

    except Exception as error:
        logger.error(
            "extractor_docx.extract_text_from_docx: "
            "Unexpected error reading DOCX '%s': %s",
            file_path,
            error,
        )
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """
    Reads a plain-text (.txt) bank statement into a raw text string.

    Some banks export a fixed-width / space-aligned plain-text statement. There is
    nothing to decode — we just read the characters — but Indian banking software
    sometimes writes non-UTF-8 files, so we try a few encodings in order before
    giving up. The returned text is handed to the SAME deterministic text parser the
    digital PDFs use (standardise_digital_pdf_transactions), so .txt statements get
    the identical header/footer/id-column handling with no extra logic.

    Parameters:
        file_path (str): Absolute path to the .txt file.

    Returns:
        str: The file's text, or "" if it cannot be read.
    """
    from pathlib import Path

    for encoding in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            with open(file_path, encoding=encoding) as f:
                text = f.read()
            logger.info(
                "extractor_docx.extract_text_from_txt: Read '%s' with encoding '%s' "
                "(%d characters).", Path(file_path).name, encoding, len(text),
            )
            return text
        except UnicodeDecodeError:
            continue
        except FileNotFoundError:
            logger.error(
                "extractor_docx.extract_text_from_txt: File not found: %s", file_path)
            return ""
        except Exception as error:
            logger.error(
                "extractor_docx.extract_text_from_txt: Error reading TXT '%s': %s",
                file_path, error)
            return ""

    logger.error(
        "extractor_docx.extract_text_from_txt: All encodings failed for '%s'.",
        file_path)
    return ""
