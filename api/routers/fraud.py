"""
fraud.py — surfaces the REAL fraud-detection engine (the 22 deterministic detectors,
suspicion scoring, clusters, Case Narrative + Prime Suspects) and the court-ready PDF
report to the React frontend.

This router ONLY calls the existing backend entry points — it never changes any detector,
threshold, score, or report logic:
    analysis  -> analysis.analysis_engine.pipeline.AnalysisPipeline(...).run()
    payload   -> reporting.report_data.build_report_context(...)   (same data the PDF uses)
    pdf       -> reporting/build_report.py under .report_venv       (same as app.py)

Both the analysis run and the PDF render are heavy, so they run in background threads and
the frontend polls a status endpoint.
"""

from __future__ import annotations

import json
import re
import sqlite3
import subprocess
import sys
import threading
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse

ROOT = Path(__file__).resolve().parents[2]  # repo root (api/routers/ -> api/ -> root)

router = APIRouter(prefix="/api/cases/{case_id}/fraud", tags=["fraud-engine"])

# In-process job state (uvicorn is one long-lived process). Keyed by (case_id, kind).
_JOBS: dict[str, dict[str, Any]] = {}
_LOCK = threading.Lock()


def _extraction_dir(case_id: str) -> Path:
    return ROOT / "outputs" / "extractions" / case_id


def _analysis_dir(case_id: str) -> Path:
    return ROOT / "analysis" / "outputs" / case_id


def _job_key(case_id: str, kind: str) -> str:
    return f"{kind}:{case_id}"


def _set_job(case_id: str, kind: str, **fields: Any) -> None:
    with _LOCK:
        _JOBS[_job_key(case_id, kind)] = {"kind": kind, **fields}


def _get_job(case_id: str, kind: str) -> dict[str, Any] | None:
    with _LOCK:
        return _JOBS.get(_job_key(case_id, kind))


# ── Deep analysis (the 22-detector engine) ───────────────────────────────────

def _run_analysis_job(case_id: str) -> None:
    try:
        from analysis.analysis_engine.config import AnalysisConfig
        from analysis.analysis_engine.pipeline import AnalysisPipeline

        AnalysisPipeline(
            input_path=_extraction_dir(case_id),
            output_dir=_analysis_dir(case_id),
            config=AnalysisConfig(),  # same config the Streamlit app uses
        ).run()
        # Drop the cached chatbot resources so the next chat load re-indexes the case WITH
        # the fresh analysis chunks (Case Narrative, suspects, findings) — this is what lets
        # the chatbot answer questions about the analysis and report.
        try:
            from api.deps import invalidate_case

            invalidate_case(case_id)
        except Exception:
            pass
        _set_job(case_id, "analysis", status="done")
    except Exception as exc:  # noqa: BLE001 — surface any failure to the UI, never crash the server
        _set_job(case_id, "analysis", status="error", error=str(exc))


@router.post("/run")
def run_analysis(case_id: str):
    """Kick off the real fraud-detection engine for a case (background)."""
    if not (_extraction_dir(case_id) / "clean_transactions.csv").exists():
        raise HTTPException(status_code=404, detail="No extraction found for this case. Run extraction first.")
    job = _get_job(case_id, "analysis")
    if job and job.get("status") == "running":
        return {"status": "running"}
    _set_job(case_id, "analysis", status="running")
    threading.Thread(target=_run_analysis_job, args=(case_id,), daemon=True).start()
    return {"status": "running"}


@router.get("/status")
def analysis_status(case_id: str):
    """Report whether deep analysis is not_run / running / done / error."""
    job = _get_job(case_id, "analysis")
    if job and job.get("status") in {"running", "error"}:
        return {"status": job["status"], "error": job.get("error", "")}
    done = (_analysis_dir(case_id) / "analysis_results.json").exists()
    return {"status": "done" if done else "not_run"}


def _build_context(case_id: str) -> dict[str, Any]:
    """Build the same holder-named, narrative payload the PDF report uses."""
    # report_data imports its sibling i18n by bare name, so reporting/ must be importable.
    reporting_dir = str(ROOT / "reporting")
    if reporting_dir not in sys.path:
        sys.path.insert(0, reporting_dir)
    from report_data import build_report_context  # type: ignore

    return build_report_context(run_id=case_id, repo_root=ROOT)


# Human names + which patterns are "priority" (the strongest flow-based evidence).
_PATTERN_NAMES = {
    1: "Duplicate Verification", 2: "Failed / Reversed Transactions", 3: "Pass-Through Routing",
    4: "Fund Pooling", 5: "Structuring / Smurfing", 7: "Circular Flow", 8: "Money Trail",
    9: "Credit-to-Cash-Out", 10: "Cross-Statement Links", 11: "Balance Parking", 12: "Hub Ranking",
    13: "Low-Value Account Testing", 14: "Reversal Clusters", 15: "Round-Value Debits",
    16: "Shared UPI Identifiers", 17: "Round-Trip Detection", 18: "Dormant Reactivation",
    19: "First-Contact Large Transfer",
}
_PRIORITY_PATTERNS = {8, 17, 7, 10}      # strongest flow-based evidence, surfaced first
_LEAD_PATTERNS = {22, 23}                 # AI / statistical leads (not scored)
_STRUCTURAL_META = {6, 21}                # graph build + ranking — never account-accusatory

_TXN_ID_RE = re.compile(r'\bacct_\d+_\d{6}\b')
_ACCT_ID_RE = re.compile(r'\bacct_\d+\b')


def _relabel_finding_text(text: str, account_map: dict[str, str]) -> str:
    """Replace internal acct_XXX_NNNNNN and acct_XXX IDs in a finding explanation
    with the real account label. Presentation only — no data is changed."""
    if not text:
        return text
    # Pass 1: txn_id patterns  (acct_001_000207 → label of the owning account)
    def _sub_txn(m: re.Match) -> str:
        tid = m.group(0)
        # Strip the 6-digit row index to get the account ID
        acc_id = "_".join(tid.rsplit("_", 1)[:-1])
        return account_map.get(acc_id, acc_id)
    text = _TXN_ID_RE.sub(_sub_txn, text)
    # Pass 2: bare account IDs  (acct_001 → Holder (number))
    def _sub_acc(m: re.Match) -> str:
        return account_map.get(m.group(0), m.group(0))
    return _ACCT_ID_RE.sub(_sub_acc, text)


def _pattern_findings(case_id: str, account_map: dict[str, str]) -> tuple[list, list]:
    """Read analysis_results.json and return (scored_pattern_findings, leads), each item
    holding only NON-EMPTY findings with holder-named accounts. Zero-finding patterns are
    dropped entirely (Priority 6). Presentation only — reads analysis output, changes nothing."""
    path = _analysis_dir(case_id) / "analysis_results.json"
    if not path.exists():
        return [], []
    data = json.loads(path.read_text(encoding="utf-8"))
    fbp = data.get("findings_by_pattern", {}) or {}

    def _label(acc: str) -> str:
        return account_map.get(str(acc), str(acc))

    scored: list = []
    leads: list = []
    for key, items in fbp.items():
        try:
            pid = int(str(key).split("_", 1)[0])
        except ValueError:
            continue
        if pid in _STRUCTURAL_META:
            continue
        accusatory = [f for f in (items or []) if f.get("accounts")]
        if not accusatory:  # Priority 6: hide zero-finding patterns
            continue
        rows = [
            {
                "accounts": [_label(a) for a in f.get("accounts", [])][:6],
                "text": _relabel_finding_text(
                    f.get("narration") or f.get("explanation", ""), account_map
                ),
                "evidence_strength": f.get("evidence_strength", ""),
            }
            for f in accusatory[:8]
        ]
        entry = {
            "pattern_id": pid,
            "name": _PATTERN_NAMES.get(pid, str(key)),
            "count": len(accusatory),
            "priority": pid in _PRIORITY_PATTERNS,
            "findings": rows,
        }
        (leads if pid in _LEAD_PATTERNS else scored).append(entry)

    scored.sort(key=lambda e: (not e["priority"], -e["count"]))
    return scored, leads


def _account_identity(case_id: str) -> dict[str, dict[str, str]]:
    """source_account_id -> {holder, account_number, label} from the extraction receipt.

    `label` prefers the REAL account number (with holder when known) and only falls
    back to the internal source id (acct_00x) when extraction captured no number.
    """
    identity: dict[str, dict[str, str]] = {}
    meta = _extraction_dir(case_id) / "metadata.json"
    if meta.exists():
        try:
            files = json.loads(meta.read_text(encoding="utf-8")).get("files", []) or []
            for f in files:
                sid = str(f.get("source_account_id", "") or "")
                if not sid:
                    continue
                ad = f.get("account_details", {}) or {}
                holder = str(ad.get("account_holder", "") or "").strip()
                number = str(ad.get("account_number", "") or "").strip()
                if number and holder:
                    label = f"{holder} ({number})"
                elif number:
                    label = number
                elif holder:
                    label = holder
                else:
                    label = sid
                identity[sid] = {"holder": holder, "account_number": number, "label": label}
        except Exception:
            pass
    return identity


def _account_map(case_id: str) -> dict[str, str]:
    """source_account_id -> 'Holder (number)' from the extraction receipt.

    Falls back to the source id only when no real number/holder was extracted, so
    the app shows genuine account numbers everywhere they are available.
    """
    return {sid: info["label"] for sid, info in _account_identity(case_id).items()}


def _resolve_label(token: str, identity: dict[str, dict[str, str]]) -> str:
    """Human label for a graph node: real number/holder if the node is an observed
    account, otherwise the token itself (a counterparty handle / UPI id / raw number)."""
    info = identity.get(str(token))
    return info["label"] if info else str(token)


@router.get("")
def get_analysis(case_id: str):
    """Return the real fraud findings: Case Narrative, Prime Suspects, ranked suspects,
    key findings and case reconstruction — exactly what the court-ready PDF contains."""
    if not (_analysis_dir(case_id) / "analysis_results.json").exists():
        raise HTTPException(status_code=404, detail="Deep analysis has not been run for this case yet.")
    try:
        ctx = _build_context(case_id)
        scored_patterns, leads = _pattern_findings(case_id, _account_map(case_id))
        raw = json.loads((_analysis_dir(case_id) / "analysis_results.json").read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Could not read analysis results: {exc}") from exc
    analysis_summary = dict(ctx.get("analysis_summary", {}) or {})
    analysis_summary["llm_status"] = raw.get("run_metadata", {}).get("llm_status", "")
    return {
        "case_narrative": ctx.get("case_narrative", {}),
        "prime_suspects": ctx.get("prime_suspects", []),
        "ranked_accounts": ctx.get("ranked_accounts", []),
        "ranked_shown": ctx.get("ranked_shown", 0),
        "ranked_total": ctx.get("ranked_total", 0),
        "analysis_summary": analysis_summary,
        "case_reconstruction": ctx.get("case_reconstruction", {}),
        "final_summary": ctx.get("final_summary", {}),
        "data_security": ctx.get("data_security", {}),
        # Per-pattern findings for the interactive Analysis UI — only patterns with real
        # findings are included (Priority 6); zero-finding patterns are dropped server-side.
        "scored_patterns": scored_patterns,
        "leads": leads,
    }


# ── Court-ready PDF report ────────────────────────────────────────────────────

def _report_pdf_path(case_id: str, language: str) -> Path:
    name = "investigation_report_kn.pdf" if language == "kn" else "investigation_report.pdf"
    return _analysis_dir(case_id) / name


def _run_report_job(case_id: str, language: str) -> None:
    venv_py = ROOT / ".report_venv" / "bin" / "python"
    out_pdf = _report_pdf_path(case_id, language)
    if not venv_py.exists():
        _set_job(case_id, "report", status="error", language=language,
                 error="Report renderer not installed (.report_venv missing). See DELIVERABLES.md setup.")
        return
    cmd = [str(venv_py), str(ROOT / "reporting" / "build_report.py"),
           "--run-id", case_id, "--ai", "--out", str(out_pdf), "--language", language, "--save-html"]
    try:
        proc = subprocess.run(cmd, cwd=str(ROOT), capture_output=True, text=True, timeout=600)
    except subprocess.TimeoutExpired:
        _set_job(case_id, "report", status="error", language=language, error="Report generation timed out.")
        return
    if out_pdf.exists():
        _set_job(case_id, "report", status="done", language=language, format="pdf")
    elif out_pdf.with_suffix(".html").exists():
        _set_job(case_id, "report", status="done", language=language, format="html")
    else:
        _set_job(case_id, "report", status="error", language=language,
                 error=(proc.stderr or proc.stdout or "unknown error")[-600:])


@router.post("/report/run")
def run_report(case_id: str, language: str = Query("en", pattern="^(en|kn)$")):
    """Generate the court-ready PDF (Case Narrative + Prime Suspects) in the background."""
    if not (_analysis_dir(case_id) / "analysis_results.json").exists():
        raise HTTPException(status_code=404, detail="Run deep analysis before generating the report.")
    job = _get_job(case_id, "report")
    if job and job.get("status") == "running":
        return {"status": "running"}
    _set_job(case_id, "report", status="running", language=language)
    threading.Thread(target=_run_report_job, args=(case_id, language), daemon=True).start()
    return {"status": "running"}


@router.get("/report/status")
def report_status(case_id: str, language: str = Query("en", pattern="^(en|kn)$")):
    job = _get_job(case_id, "report")
    if job and job.get("status") in {"running", "error"}:
        return {"status": job["status"], "error": job.get("error", ""), "language": job.get("language", language)}
    pdf = _report_pdf_path(case_id, language)
    if pdf.exists():
        return {"status": "done", "format": "pdf", "language": language}
    if pdf.with_suffix(".html").exists():
        return {"status": "done", "format": "html", "language": language}
    return {"status": "not_run", "language": language}


@router.get("/report/download")
def download_report(case_id: str, language: str = Query("en", pattern="^(en|kn)$")):
    pdf = _report_pdf_path(case_id, language)
    if pdf.exists():
        return FileResponse(pdf, media_type="application/pdf", filename=pdf.name)
    html = pdf.with_suffix(".html")
    if html.exists():  # graceful fallback when Chromium/PDF is unavailable
        return FileResponse(html, media_type="text/html", filename=html.name)
    raise HTTPException(status_code=404, detail="No report generated yet for this case/language.")


# ── Interactive investigation graphs ─────────────────────────────────────────
# The analysis run already emits normalized UI graph JSON under
# <analysis_dir>/graphs/ui_graphs/ (see analysis_engine/graph_generator.py). This
# serves those files as-is — SQL stays the source of truth, this is a read-only
# projection. The contract is backend-agnostic ({nodes,edges}/{flows}/{accounts}/
# {trails}) so a future Neo4j-backed builder can swap in without a frontend change.
_UI_GRAPHS: dict[str, str] = {
    "money_flow": "money_flow_network_3d.json",
    "balance": "balance_graph_all_accounts.json",
    "money_trail": "money_trail_all_accounts.json",
    "sankey": "sankey_flow.json",
}


@router.get("/graphs")
def list_graphs(case_id: str):
    """Which interactive graphs this case produced (best-effort, may be partial)."""
    ui_dir = _analysis_dir(case_id) / "graphs" / "ui_graphs"
    db_exists = (_analysis_dir(case_id) / "analysis.db").exists()
    available = {name: (ui_dir / fname).exists() for name, fname in _UI_GRAPHS.items()}
    # Sankey and Money Trail are computed on demand from analysis.db (not pre-generated
    # files), so their availability tracks the database, not a file on disk.
    available["sankey"] = db_exists
    available["money_trail"] = db_exists
    return {"case_id": case_id, "available": available}


def _label_graph_nodes(payload: dict[str, Any], identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """Attach real account number/holder to every node in a money-flow payload.

    Observed accounts (acct_00x) resolve to their REAL number/holder from the
    extraction receipt; counterparties keep their handle. Purely presentational —
    the underlying id is preserved so every element still traces back to SQL.
    """
    data = payload.get("data")
    if not isinstance(data, dict):
        return payload
    for node in data.get("nodes", []) or []:
        if not isinstance(node, dict):
            continue
        nid = str(node.get("id", ""))
        info = identity.get(nid)
        node["label"] = info["label"] if info else nid
        node["account_number"] = info["account_number"] if info else ""
        node["holder"] = info["holder"] if info else ""
    return payload


# Transaction-type Sankey — same categorisation the chatbot uses (chatbot/graph_viz.py
# + investigation_rules.py), rebuilt here from analysis.db so it works on any existing
# case without re-running analysis. 3 layers: account -> txn-type bucket -> External.
_TXN_TYPE_PATTERNS: list[tuple[str, re.Pattern]] = [
    ("UPI", re.compile(r"\bUPI\b", re.IGNORECASE)),
    ("NEFT", re.compile(r"\bNEFT\b", re.IGNORECASE)),
    ("RTGS", re.compile(r"RTGS", re.IGNORECASE)),  # "INRTGS" prefix also valid
    ("IMPS", re.compile(r"\bIMPS\b", re.IGNORECASE)),
    ("Cheque", re.compile(r"\b(cheque|chq|cheq)\b", re.IGNORECASE)),
    ("Cash", re.compile(r"\b(?:cash|atm|withdrawal|deposit|cdm|cwdl)\b", re.IGNORECASE)),
]


def _classify_txn_type(narration: str) -> str:
    text = str(narration or "")
    for label, pattern in _TXN_TYPE_PATTERNS:
        if pattern.search(text):
            return label
    return "Other"


# The mode/channel code a bank prints at the very START of a narration (BLKRTGS, BLKIFT,
# UPI, NEFT, IMPS-OPW, ATM-NFS, NACH, ECS …). For the Money Trail badges we show this
# VERBATIM rather than bucketing it, so "BLKRTGS" reads as BLKRTGS (not RTGS) and "BLKIFT"
# as BLKIFT (not OTHER). Rule (generic, no bank/format hardcoding): the leading run of
# letters, extended across a hyphen only when the next segment is also ≥2 letters — so a
# real sub-mode like IMPS-OPW / FI-AEPS-OFFUS-ISS-CW is kept, while a reference tail like
# NEFT-N49652837781 or CASHKIOSK-2653L is cut at the digits. Stops at the first '/',
# space, or digit.
_LEADING_MODE_RE = re.compile(r"[A-Za-z]+(?:-[A-Za-z]{2,})*")


def _leading_narration_token(narration: str) -> str:
    """The transaction mode/channel token as printed at the start of the narration."""
    text = str(narration or "").strip()
    if not text:
        return "Other"
    match = _LEADING_MODE_RE.match(text)
    if match:
        return match.group(0).upper()
    # Narration begins with a reference number / symbol, not a mode code (e.g.
    # "913628731289:Int.Pd:…"). A raw numeric prefix is meaningless as a badge, so fall
    # back to the rail classifier, which scans the whole narration for a known channel
    # and otherwise returns "Other".
    return _classify_txn_type(text)


_SKIP_PARTS = frozenset({"upi", "neft", "rtgs", "imps", "payment", "dr", "cr", "by", "to", "transfer"})
_UTR_RE = re.compile(r"^[A-Z]{3,6}\d{8,}$")           # e.g. HDFCN2506110993676
_TIME_RE = re.compile(r"^\d{1,2}:\d{2}(:\d{2})?$")
_DATE_SUFFIX_RE = re.compile(r"\d{2}-[A-Z][a-z]{2}-\d{2}")  # 28-Apr-25 style date → processing suffix
_NEFT_RTGS_NAME_RE = re.compile(r"(?:NEFT|RTGS|IMPS)[/-]?(?:[A-Z0-9]{6,})[/-]([A-Za-z][A-Za-z\s\.&]+)", re.IGNORECASE)


def _extract_sender(narration: str, counterparty_account: Any, counterparty_name_raw: Any) -> str:
    """Deterministically extract sender name from narration + counterparty fields. No LLM."""
    def _clean(v: Any) -> str:
        s = str(v or "").strip()
        return "" if s.lower() in ("nan", "none", "null", "<na>", "") else s

    name_raw = _clean(counterparty_name_raw)
    if name_raw:
        return name_raw

    text = _clean(narration)

    # NEFT/RTGS: "NEFT-UTR-SENDER NAME" or "RTGS:UTR/SENDER NAME"
    m = _NEFT_RTGS_NAME_RE.search(text)
    if m:
        name = m.group(1).strip()
        if len(name) > 2:
            return name

    # Split on "/" and take the last meaningful token
    parts = [p.strip() for p in text.split("/") if p.strip()]
    for part in reversed(parts):
        if _TIME_RE.match(part):
            continue
        if re.match(r"^\d+$", part):
            continue
        if _UTR_RE.match(part):
            continue
        if _DATE_SUFFIX_RE.search(part):          # e.g. "P2AMOB TRF 28-Apr-25"
            continue
        if part.lower() in _SKIP_PARTS:
            continue
        if "@" in part:
            # UPI ID — prefer counterparty_account if it's a UPI handle
            cp = _clean(counterparty_account)
            if cp and "@" in cp:
                return cp
            return part
        return part

    cp = _clean(counterparty_account)
    if cp and not cp.replace(" ", "").isdigit():
        return cp

    return ""


def _enrich_money_trail(payload: dict[str, Any], db_path: Path, identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """Add source_credit_details (date/type/sender/reference) and allocation txn_type from analysis.db."""
    trails = payload.get("trails") or []
    if not trails or not db_path.exists():
        return payload

    # Collect all source credit txn_ids that need SQL look-up
    src_ids = [t.get("source_credit_txn_id") for t in trails if t.get("source_credit_txn_id")]
    if not src_ids:
        return payload

    con = sqlite3.connect(str(db_path))
    try:
        placeholders = ",".join("?" * len(src_ids))
        rows = con.execute(
            f"SELECT txn_id, date, time, narration, reference, counterparty_account, counterparty_name_raw "
            f"FROM transactions WHERE txn_id IN ({placeholders})",
            src_ids,
        ).fetchall()
    except sqlite3.Error:
        return payload
    finally:
        con.close()

    txn_map = {
        r[0]: {"date": r[1] or "", "time": r[2] or "", "narration": r[3] or "",
                "reference": r[4] or "", "counterparty_account": r[5], "counterparty_name_raw": r[6]}
        for r in rows
    }
    observed_numbers: set[str] = {
        info["account_number"] for info in identity.values() if info.get("account_number")
    }

    for trail in trails:
        src_id = trail.get("source_credit_txn_id")
        if src_id and src_id in txn_map:
            row = txn_map[src_id]
            nar = row["narration"]
            txn_type = _leading_narration_token(nar)
            sender = _extract_sender(nar, row["counterparty_account"], row["counterparty_name_raw"])
            # Flag if sender is an observed account within the investigation
            cp_acc = str(row["counterparty_account"] or "").strip()
            is_observed = cp_acc in observed_numbers or any(
                cp_acc == info.get("account_number") for info in identity.values()
            )
            trail["source_credit_details"] = {
                "date": row["date"],
                "time": row["time"],
                "narration": nar,
                "reference": row["reference"],
                "txn_type": txn_type,
                "sender": sender,
                "is_observed_account": is_observed,
            }

        # Classify txn_type for each allocation from its narration
        for alloc in trail.get("allocations") or []:
            alloc["txn_type"] = _leading_narration_token(str(alloc.get("narration") or ""))


# ── Money Trail — EVERY incoming credit, LAZY per-credit tracing ──────────────
# The fraud detector only auto-traces threshold-gated (high-value) credits, which is
# right for scoring but too narrow for an investigator who must be able to open the
# trail behind ANY incoming credit.
#
# Tracing is O(1 SQL read of the account history) PER credit, so eagerly tracing every
# credit does not scale: real cases carry 10k+ credits (~11 min to trace all). Instead
# we split the work:
#   • the INDEX (_build_money_trail_index) lists every credit grouped by account,
#     highest→lowest, with NO tracing — one SELECT, instant even at 10k+ credits;
#   • the TRAIL (_trace_one_credit) runs the real FIFO algorithm for a SINGLE credit,
#     on demand, only when the investigator opens it (~tens of ms).
# The FIFO tracing logic itself is reused untouched from analysis_engine.
_MONEY_TRAIL_INDEX_FILE = "money_trail_credit_index.json"
# Bump when the payload SHAPE or any derived field (e.g. the txn_type badge rule) changes,
# so a stale cache from an older build is recomputed instead of served.
_MONEY_TRAIL_CACHE_VERSION = 4
# Small in-process cache of already-traced single credits, keyed by (case_id, txn_id),
# so re-opening a credit (or toggling back to it) is instant. Bounded to avoid growth.
_TRAIL_CACHE: dict[tuple[str, str], dict[str, Any]] = {}
_TRAIL_CACHE_MAX = 512


def _build_money_trail_index(db_path: Path, identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """List every eligible incoming credit grouped by account, highest→lowest — WITHOUT
    tracing (that is done lazily per selected credit). One SELECT; scales to 10k+ credits.

    Each credit carries the summary the Level-2 list needs (amount, date, narration,
    sender, txn_type badge, reference, whether the sender is an observed account). The
    FIFO trail (allocations / trace status) is fetched on demand via _trace_one_credit."""
    from analysis.analysis_engine.config import AnalysisConfig

    epsilon = AnalysisConfig().money_epsilon
    observed_numbers = {info["account_number"] for info in identity.values() if info.get("account_number")}
    con = sqlite3.connect(str(db_path))
    con.row_factory = sqlite3.Row
    grouped: dict[str, list[dict[str, Any]]] = {}
    try:
        rows = con.execute(
            "SELECT txn_id, account_id, credit_amount, date, time, narration, reference, "
            "counterparty_account, counterparty_name_raw FROM transactions "
            "WHERE eligible_for_detection = 1 AND credit_amount > ? AND date IS NOT NULL "
            "ORDER BY account_id, credit_amount DESC, source_order, row_id",
            (epsilon,),
        ).fetchall()
    finally:
        con.close()

    for row in rows:
        narration = row["narration"] or ""
        cp_acc = str(row["counterparty_account"] or "").strip()
        grouped.setdefault(str(row["account_id"]), []).append(
            {
                "txn_id": str(row["txn_id"]),
                "credited_amount": float(row["credit_amount"] or 0.0),
                "date": row["date"] or "",
                "time": row["time"] or "",
                "narration": narration,
                "reference": row["reference"] or "",
                "txn_type": _leading_narration_token(narration),
                "sender": _extract_sender(narration, row["counterparty_account"], row["counterparty_name_raw"]),
                "is_observed_account": bool(cp_acc) and cp_acc in observed_numbers,
            }
        )

    accounts = [
        {
            "account_id": account_id,
            "credit_count": len(credits),
            "total_credited": round(sum(c["credited_amount"] for c in credits), 2),
            "credits": credits,  # already highest→lowest within the account (query order)
        }
        for account_id, credits in grouped.items()
    ]
    accounts.sort(key=lambda a: a["total_credited"], reverse=True)
    return {"type": "money_trail_by_account", "accounts": accounts}


def _money_trail_by_account(case_id: str, identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """Serve the credit INDEX (Level 1 + Level 2), computing + caching it on first use.

    The index does no tracing, so it is fast; it is still cached to disk (keyed on the
    analysis.db mtime + version) to avoid re-running the per-credit sender/type parsing."""
    db_path = _analysis_dir(case_id) / "analysis.db"
    if not db_path.exists():
        raise HTTPException(status_code=404, detail="Analysis has not been run for this case yet.")

    cache = _analysis_dir(case_id) / "graphs" / "ui_graphs" / _MONEY_TRAIL_INDEX_FILE
    payload: dict[str, Any] | None = None
    try:
        if cache.exists() and cache.stat().st_mtime >= db_path.stat().st_mtime:
            cached = json.loads(cache.read_text(encoding="utf-8"))
            if int(cached.get("cache_version", 0)) == _MONEY_TRAIL_CACHE_VERSION:
                payload = cached
    except Exception:
        payload = None

    if payload is None:
        payload = _build_money_trail_index(db_path, identity)
        payload["cache_version"] = _MONEY_TRAIL_CACHE_VERSION
        try:
            cache.parent.mkdir(parents=True, exist_ok=True)
            cache.write_text(json.dumps(payload), encoding="utf-8")
        except Exception:
            pass  # a failed cache write only costs a recompute next time — never fatal

    # Presentation labels (id -> real account number/holder), attached fresh & cheaply.
    payload["labels"] = {sid: info["label"] for sid, info in identity.items()}
    for acc in payload.get("accounts", []):
        info = identity.get(str(acc.get("account_id", "")))
        acc["label"] = info["label"] if info else str(acc.get("account_id", ""))
        acc["account_number"] = info["account_number"] if info else ""
        acc["holder"] = info["holder"] if info else ""
    return payload


def _trace_one_credit(case_id: str, txn_id: str, identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """Run the real FIFO money-trail for a SINGLE credit, on demand (Level 3).

    Reuses analysis_engine.detectors.money_trail.trace_credit unchanged, then enriches
    the trail the same way the index enriches summaries. Cached in-process per credit."""
    key = (case_id, txn_id)
    hit = _TRAIL_CACHE.get(key)
    if hit is not None:
        return hit

    db_path = _analysis_dir(case_id) / "analysis.db"
    if not db_path.exists():
        raise HTTPException(status_code=404, detail="Analysis has not been run for this case yet.")

    from analysis.analysis_engine.config import AnalysisConfig
    from analysis.analysis_engine.detectors.money_trail import trace_credit

    con = sqlite3.connect(str(db_path))
    con.row_factory = sqlite3.Row
    try:
        finding = trace_credit(
            con,
            str(txn_id),
            AnalysisConfig(),
            trigger_reason="manual_view",
            gated_reason="investigator_selected_credit",
        )
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    finally:
        con.close()

    details = finding.details or {}
    trail = {
        "finding_id": finding.finding_id,
        "accounts": list(finding.accounts),
        "source_credit_txn_id": details.get("source_credit_txn_id"),
        "credited_amount": details.get("credited_amount"),
        "trace_status": details.get("trace_status"),
        "allocations": details.get("allocations", []),
    }
    _enrich_money_trail({"trails": [trail]}, db_path, identity)

    if len(_TRAIL_CACHE) >= _TRAIL_CACHE_MAX:
        _TRAIL_CACHE.pop(next(iter(_TRAIL_CACHE)))  # drop oldest (FIFO) to bound memory
    _TRAIL_CACHE[key] = trail
    return trail


# ── Money Trail → Word (.docx) export in the ORIGINAL statement layout ────────
# The export reproduces the uploaded statement's ORIGINAL column names + order (the
# display schema captured at extraction), NOT our internal schema — and populates it
# with only the transactions of the selected trail (the source credit + the debits
# that consume it). Balance is excluded by default; columns we cannot populate are
# dropped so the table stays clean. Everything is read from artefacts already on disk
# (metadata.json schema + analysis.db values), so there is no re-reading of uploads.


def _indian_amount(value: Any) -> str:
    """Format a money value with Indian digit grouping (e.g. 12,34,567.00). Blank for 0."""
    try:
        num = float(str(value).replace(",", "").strip() or 0)
    except (TypeError, ValueError):
        return ""
    if abs(num) < 0.005:
        return ""
    whole, dec = f"{num:.2f}".split(".")
    sign = "-" if whole.startswith("-") else ""
    whole = whole.lstrip("-")
    if len(whole) > 3:
        head, tail = whole[:-3], whole[-3:]
        head = re.sub(r"(\d)(?=(\d\d)+$)", r"\1,", head)
        grouped = f"{head},{tail}"
    else:
        grouped = whole
    return f"{sign}{grouped}.{dec}"


def _cell_value(field: str, payload: dict[str, Any]) -> str:
    """Render one original-layout cell from a transaction's normalized values."""
    def g(k: str) -> str:
        v = payload.get(k, "")
        s = str(v).strip()
        return "" if s.lower() in ("nan", "none", "null", "<na>") else s

    if field == "Date":
        return g("Date")
    if field == "Time":
        return g("Time")
    if field == "Date_Time":
        return f"{g('Date')} {g('Time')}".strip()
    if field == "Narration":
        return g("Narration")
    if field == "Debit":
        return _indian_amount(g("Debit"))
    if field == "Credit":
        return _indian_amount(g("Credit"))
    if field == "Balance":
        return _indian_amount(g("Balance"))
    if field == "Cheque_Number":
        return g("Cheque_Number")
    if field == "Transaction_Reference":
        return g("Transaction_Reference") or g("Reference_Number")
    if field == "ref_or_cheque":
        return g("Cheque_Number") or g("Transaction_Reference") or g("Reference_Number")
    if field == "Amount":
        return _indian_amount(g("Debit")) or _indian_amount(g("Credit"))
    return ""


def _load_display_schema(case_id: str, account_id: str) -> list[dict[str, str]]:
    """The account's original column layout captured at extraction (metadata.json).
    Falls back to a generic layout for older runs that predate schema capture."""
    meta = _extraction_dir(case_id) / "metadata.json"
    if meta.exists():
        try:
            files = json.loads(meta.read_text(encoding="utf-8")).get("files", []) or []
            for f in files:
                if str(f.get("source_account_id", "")) == str(account_id):
                    cols = ((f.get("display_schema") or {}).get("columns")) or []
                    if cols:
                        return cols
        except Exception:
            pass
    # Fallback (pre-capture runs): a neutral, generic layout.
    return [
        {"name": "Date", "field": "Date"},
        {"name": "Narration", "field": "Narration"},
        {"name": "Debit", "field": "Debit"},
        {"name": "Credit", "field": "Credit"},
    ]


def _trail_payloads(db_path: Path, txn_ids: list[str]) -> dict[str, dict[str, Any]]:
    """Original per-transaction normalized values (raw_payload_json) keyed by txn_id."""
    if not txn_ids:
        return {}
    con = sqlite3.connect(str(db_path))
    try:
        placeholders = ",".join("?" * len(txn_ids))
        rows = con.execute(
            f"SELECT txn_id, raw_payload_json, date, time, narration, debit_amount, "
            f"credit_amount, balance, reference FROM transactions WHERE txn_id IN ({placeholders})",
            txn_ids,
        ).fetchall()
    except sqlite3.Error:
        return {}
    finally:
        con.close()
    out: dict[str, dict[str, Any]] = {}
    for r in rows:
        payload: dict[str, Any] = {}
        if r[1]:
            try:
                payload = json.loads(r[1])
            except Exception:
                payload = {}
        # Backfill from typed columns if raw_payload lacked a key (older extractions).
        payload.setdefault("Date", r[2] or "")
        payload.setdefault("Time", r[3] or "")
        payload.setdefault("Narration", r[4] or "")
        payload.setdefault("Debit", r[5] if r[5] is not None else "")
        payload.setdefault("Credit", r[6] if r[6] is not None else "")
        payload.setdefault("Balance", r[7] if r[7] is not None else "")
        payload.setdefault("Transaction_Reference", r[8] or "")
        out[str(r[0])] = payload
    return out


def _build_trail_docx(case_id: str, txn_id: str, identity: dict[str, dict[str, str]]) -> tuple[bytes, str]:
    """Build the .docx for one credit's money trail. Returns (bytes, filename)."""
    from docx import Document
    from docx.shared import Pt

    trail = _trace_one_credit(case_id, txn_id, identity)
    account_id = (trail.get("accounts") or [""])[0]
    # Transactions of this trail: the source credit + each consuming debit, oldest first.
    ordered_ids = [str(trail.get("source_credit_txn_id") or txn_id)]
    ordered_ids += [str(a.get("debit_txn_id")) for a in (trail.get("allocations") or []) if a.get("debit_txn_id")]

    db_path = _analysis_dir(case_id) / "analysis.db"
    payloads = _trail_payloads(db_path, ordered_ids)
    rows = [payloads[i] for i in ordered_ids if i in payloads]
    rows.sort(key=lambda p: (str(p.get("Date", "")), str(p.get("Time", ""))))

    # Columns: original layout, excluding Balance and any column we cannot populate.
    schema = _load_display_schema(case_id, account_id)
    columns = [c for c in schema if c.get("field") and c.get("field") != "Balance"]
    if not columns:  # extreme fallback so a document is always produced
        columns = [{"name": "Date", "field": "Date"}, {"name": "Narration", "field": "Narration"},
                   {"name": "Debit", "field": "Debit"}, {"name": "Credit", "field": "Credit"}]

    info = identity.get(str(account_id)) or {}
    acct_label = info.get("label") or str(account_id)
    src = trail.get("source_credit_details") or {}

    doc = Document()
    doc.add_heading("Money Trail — Transaction Export", level=1)
    meta_p = doc.add_paragraph()
    meta_p.add_run("Account: ").bold = True
    meta_p.add_run(f"{acct_label}\n")
    meta_p.add_run("Traced credit: ").bold = True
    meta_p.add_run(f"{_indian_amount(trail.get('credited_amount'))} on {src.get('date','')} "
                   f"{src.get('time','')}".strip() + "\n")
    meta_p.add_run("Transactions in this trail: ").bold = True
    meta_p.add_run(str(len(rows)))

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for j, col in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col.get("name", ""))
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for payload in rows:
        cells = table.add_row().cells
        for j, col in enumerate(columns):
            cells[j].text = _cell_value(col.get("field", ""), payload)
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    safe_acct = re.sub(r"[^A-Za-z0-9_-]+", "_", str(info.get("account_number") or account_id))
    filename = f"money_trail_{safe_acct}_{txn_id}.docx"
    return buf.getvalue(), filename


# Counterparty identity extraction from narration — deterministic, no LLM. Every eligible
# transaction gets a counterparty node even when the resolver left counterparty_account empty,
# because the name is usually sitting in the narration (e.g. "UPI:PAY:123/NIDHI PILLAI /Kotak").
_CP_NAME_PATTERNS: list[re.Pattern] = [
    re.compile(r"(?:UPI|IMPS)[:/](?:PAY|REC|DR|CR)[:/]\d+/\s*([A-Za-z][A-Za-z\s\.]{1,29}?)\s*/", re.IGNORECASE),
    re.compile(r"(?:NEFT|RTGS|IMPS)[/:-]?[A-Z0-9]{6,}[/-]\s*([A-Za-z][A-Za-z\s\.&]{1,29}?)(?:\s+\d|/|$)", re.IGNORECASE),
    re.compile(r"(?:TRF\s+frm|frm|FROM)\s+([A-Za-z][A-Za-z\s]{1,29}?)(?:\s+\d|\s*/|$)", re.IGNORECASE),
    re.compile(r"INET-IMPS-(?:CR|DR)/\s*([A-Za-z][A-Za-z\s\.]{1,29}?)\s*/", re.IGNORECASE),
]
# Narration → non-person category (bank-internal / cash / tax), so charges cluster together.
_CP_CATEGORIES: list[tuple[str, re.Pattern]] = [
    ("Bank Charges", re.compile(r"\b(charge|fee|fees|nmmb|slabwise|amc|penal|min\s*bal)\b", re.IGNORECASE)),
    ("Tax / GST", re.compile(r"\b(gst|cgst|sgst|igst|tds|tax)\b", re.IGNORECASE)),
    ("ATM / Cash", re.compile(r"\b(atm|nfscwd|cash-p|cwdl|cash\s*wdl|self)\b", re.IGNORECASE)),
    ("Interest", re.compile(r"\b(int\.?pd|interest|int\s*credit)\b", re.IGNORECASE)),
]
_BANK_HINT_RE = re.compile(r"bank|karb|kkbk|hdfc|icici|axis|sbin|union|canara|federal|kotak|ybl|okaxis|oksbi|okicici|paytm|ibl|utib|sibl|fbl|iib", re.IGNORECASE)


def _extract_counterparty(
    narration: str,
    counterparty_account: Any,
    counterparty_name_raw: Any,
    identity: dict[str, dict[str, str]],
    observed_sids: set[str],
) -> tuple[str, str]:
    """Return (label, kind) for the other party of a transaction. kind is one of
    observed | person | upi | account | charge | unknown. Deterministic; no LLM."""
    def _clean(v: Any) -> str:
        s = str(v or "").strip()
        return "" if s.lower() in ("nan", "none", "null", "<na>", "") else s

    cp = _clean(counterparty_account)
    if cp:
        if cp in observed_sids:
            info = identity.get(cp)
            return (info["label"] if info else cp), "observed"
        name = _clean(counterparty_name_raw)
        if "@" in cp:
            return (name.title() if name else cp), "upi"
        if cp.replace(" ", "").isdigit() and 8 <= len(cp.replace(" ", "")) <= 20:
            return (f"{name.title()} ({cp})" if name else cp), "account"
        return (name.title() if name else cp), "account"

    # Unresolved → mine the narration.
    text = _clean(narration)
    name_raw = _clean(counterparty_name_raw)

    # Bank-internal categories (ATM / charges / tax / interest) are checked BEFORE any name
    # so an "ATM Cash-..." narration (or a raw name of literally "ATM") never reads as a person.
    for label, pattern in _CP_CATEGORIES:
        if pattern.search(text) or (name_raw and pattern.search(name_raw)):
            return label, "charge"

    if name_raw and not _BANK_HINT_RE.fullmatch(name_raw):
        return name_raw.title(), "person"

    for pattern in _CP_NAME_PATTERNS:
        m = pattern.search(text)
        if m:
            name = re.sub(r"\s+", " ", m.group(1)).strip(" .")
            if len(name) >= 3 and not _BANK_HINT_RE.fullmatch(name):
                return name.title(), "person"

    return "Unidentified", "unknown"


def _full_transaction_graph(case_id: str, identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """Every eligible transaction as a directed graph — the complete money-flow knowledge graph.

    Each observed account is a hub; every transaction becomes an edge to a counterparty node.
    Counterparties are resolved from counterparty_account when available, otherwise mined from the
    narration (name / bank-charge category), so NO transaction is dropped for lack of a resolver hit.
    Edges are aggregated per (account, counterparty, direction) but keep txn_count + total amount, so
    all transactions are represented. Inter-observed credits are skipped (the sender's debit covers them).
    """
    db = _analysis_dir(case_id) / "analysis.db"
    empty = {"type": "full_transaction_graph", "nodes": [], "edges": [],
             "total_transactions": 0, "resolved_transactions": 0, "accounts": []}
    if not db.exists():
        return empty

    con = sqlite3.connect(str(db))
    try:
        total = con.execute(
            "SELECT COUNT(*) FROM transactions WHERE eligible_for_detection=1"
        ).fetchone()[0]
        observed_sids = {r[0] for r in con.execute(
            "SELECT DISTINCT source_account_id FROM transactions WHERE eligible_for_detection=1"
        ).fetchall()}
        rows = con.execute(
            "SELECT source_account_id, counterparty_account, counterparty_name_raw, "
            "debit_amount, credit_amount, narration, date, txn_id "
            "FROM transactions WHERE eligible_for_detection=1 ORDER BY date"
        ).fetchall()
    except sqlite3.Error:
        return empty
    finally:
        con.close()

    def _obs_label(sid: str) -> str:
        info = identity.get(sid)
        return info["label"] if info else sid

    nodes: dict[str, dict[str, Any]] = {}
    # Pre-add observed accounts so they're always present as hubs
    for sid in observed_sids:
        nodes[sid] = {"id": sid, "label": _obs_label(sid), "kind": "observed",
                      "total_volume": 0.0, "txn_count": 0}

    # Node id for a mined counterparty: keep it unique per (label, kind) so identical
    # names collapse to one hub and shared parties link multiple accounts together.
    def _cp_node_id(label: str, kind: str) -> str:
        return f"cp::{kind}::{label.lower()}"

    edge_agg: dict[tuple[str, str], dict[str, Any]] = {}
    counted = 0

    for src_sid, cp_acc, cp_name, debit, credit, narration, date, txn_id in rows:
        debit = float(debit or 0)
        credit = float(credit or 0)
        if debit <= 0 and credit <= 0:
            continue
        label, kind = _extract_counterparty(narration, cp_acc, cp_name, identity, observed_sids)
        counted += 1

        if kind == "observed":
            # Inter-observed: only record from the debit side to avoid duplicate edges.
            cp_sid = next((sid for sid in observed_sids if _obs_label(sid) == label), None)
            if credit > 0:
                continue
            cp_node = cp_sid or _cp_node_id(label, kind)
        elif kind == "unknown":
            # Truly unidentifiable → a per-transaction leaf, so they spray out as
            # individual endpoints instead of collapsing into one misleading mega-hub.
            cp_node = f"cp::unknown::{txn_id}"
        else:
            cp_node = _cp_node_id(label, kind)

        if debit > 0:
            s, t, amount = src_sid, cp_node, debit
        else:
            s, t, amount = cp_node, src_sid, credit
        if amount <= 0:
            continue

        for nid, nlabel, nkind in ((s, label, kind), (t, label, kind)):
            if nid not in nodes:
                nodes[nid] = {"id": nid, "label": nlabel, "kind": nkind,
                              "total_volume": 0.0, "txn_count": 0}
        nodes[s]["total_volume"] += amount
        nodes[t]["total_volume"] += amount
        nodes[src_sid]["txn_count"] += 1

        key = (s, t)
        if key not in edge_agg:
            edge_agg[key] = {
                "source": s, "target": t, "amount": 0.0, "txn_count": 0,
                "first_date": date or "", "last_date": date or "",
                "txn_type": _classify_txn_type(str(narration or "")),
                "sample_narration": str(narration or "")[:120],
            }
        bucket = edge_agg[key]
        bucket["amount"] += amount
        bucket["txn_count"] += 1
        if date:
            if not bucket["first_date"] or date < bucket["first_date"]:
                bucket["first_date"] = date
            if date > bucket["last_date"]:
                bucket["last_date"] = date

    # Account roster for the frontend's per-account toggles (id, label, txn_count)
    accounts = [
        {"id": sid, "label": _obs_label(sid), "txn_count": nodes[sid]["txn_count"]}
        for sid in sorted(observed_sids)
    ]

    return {
        "type": "full_transaction_graph",
        "total_transactions": total,
        "resolved_transactions": counted,
        "accounts": accounts,
        "nodes": list(nodes.values()),
        "edges": list(edge_agg.values()),
    }


def _transaction_type_sankey(case_id: str, identity: dict[str, dict[str, str]]) -> dict[str, Any]:
    """3-layer Sankey (account -> transaction-type bucket -> External) built from the
    run's transactions. Accounts show real numbers; link value = summed amount."""
    db = _analysis_dir(case_id) / "analysis.db"
    empty = {"type": "sankey_transaction_type", "nodes": [], "links": []}
    if not db.exists():
        return empty
    con = sqlite3.connect(str(db))
    try:
        rows = con.execute(
            "SELECT source_account_id, account_holder, account_number, narration, "
            "debit_amount, credit_amount FROM transactions"
        ).fetchall()
    except sqlite3.Error:
        return empty
    finally:
        con.close()

    nodes: list[dict[str, str]] = []
    node_index: dict[str, int] = {}

    def _idx(label: str, kind: str) -> int:
        if label not in node_index:
            node_index[label] = len(nodes)
            nodes.append({"label": label, "kind": kind})
        return node_index[label]

    def _label(sid: str, holder: str, number: str) -> str:
        info = identity.get(str(sid))
        if info and info.get("label"):
            return info["label"]
        holder, number = (holder or "").strip(), (number or "").strip()
        if number and holder:
            return f"{holder} ({number})"
        return number or holder or f"Acct {sid}"

    link_totals: dict[tuple[int, int], float] = {}
    for sid, holder, number, narration, deb, cred in rows:
        deb = float(deb or 0.0)
        cred = float(cred or 0.0)
        if deb > 0:
            direction, amount = "debit", deb
        elif cred > 0:
            direction, amount = "credit", cred
        else:
            continue
        account_label = _label(sid, holder, number)
        bucket = _classify_txn_type(narration)
        if direction == "debit":
            s, b, d = _idx(account_label, "account"), _idx(bucket, "bucket"), _idx("External", "external")
        else:
            s, b, d = _idx("External", "external"), _idx(bucket, "bucket"), _idx(account_label, "account")
        link_totals[(s, b)] = link_totals.get((s, b), 0.0) + amount
        link_totals[(b, d)] = link_totals.get((b, d), 0.0) + amount

    links = [{"source": s, "target": t, "value": round(v, 2)} for (s, t), v in link_totals.items()]
    return {"type": "sankey_transaction_type", "nodes": nodes, "links": links}


@router.get("/graphs/money_trail/credit/{txn_id}")
def get_money_trail_credit(case_id: str, txn_id: str):
    """FIFO money trail for ONE incoming credit, traced on demand (Level 3 of the UI).

    Declared before the generic /graphs/{name} route so it is matched first. Keeping the
    trace lazy (one credit per request) is what lets the Money Trail scale to cases with
    tens of thousands of credits — the Level-2 list loads instantly and only the opened
    credit is traced."""
    return _trace_one_credit(case_id, txn_id, _account_identity(case_id))


@router.get("/graphs/money_trail/credit/{txn_id}/docx")
def download_money_trail_docx(case_id: str, txn_id: str):
    """Download this credit's money trail as a Word (.docx) document in the ORIGINAL
    statement layout (original column names + order, Balance excluded)."""
    from fastapi.responses import Response

    content, filename = _build_trail_docx(case_id, txn_id, _account_identity(case_id))
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/graphs/{name}")
def get_graph(case_id: str, name: str):
    """Return one investigation graph's JSON payload for interactive rendering."""
    # Dynamically-computed graphs (not pre-generated files)
    if name == "all_flows":
        if not (_analysis_dir(case_id) / "analysis.db").exists():
            raise HTTPException(status_code=404, detail="Analysis has not been run for this case yet.")
        return _full_transaction_graph(case_id, _account_identity(case_id))

    fname = _UI_GRAPHS.get(name)
    if fname is None:
        raise HTTPException(status_code=404, detail=f"Unknown graph '{name}'.")
    # Sankey is computed on demand (transaction-type view) from analysis.db, so it works
    # for any case whose analysis has run — no dependency on a pre-generated file.
    if name == "sankey":
        if not (_analysis_dir(case_id) / "analysis.db").exists():
            raise HTTPException(status_code=404, detail="Analysis has not been run for this case yet.")
        return _transaction_type_sankey(case_id, _account_identity(case_id))
    # Money Trail is computed on demand from analysis.db and covers EVERY incoming credit
    # (grouped by account, highest→lowest), not just the threshold-gated ones — so an
    # investigator can open the trail behind any credit. Result is cached to disk.
    if name == "money_trail":
        return _money_trail_by_account(case_id, _account_identity(case_id))
    path = _analysis_dir(case_id) / "graphs" / "ui_graphs" / fname
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Graph '{name}' has not been generated for this case yet.")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
        identity = _account_identity(case_id)
        # Universal id -> real-label map so every graph can resolve acct_00x to the
        # real account number/holder client-side (counterparties fall through to their id).
        payload["labels"] = {sid: info["label"] for sid, info in identity.items()}
        if name == "money_flow":
            _label_graph_nodes(payload, identity)
        return payload
    except Exception:  # never fail the graph over labeling — fall back to raw file
        return FileResponse(path, media_type="application/json")
